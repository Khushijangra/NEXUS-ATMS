from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "results" / "DTI_Project_Final_Report_NEXUS_ATMS.docx"
OUT_MD = ROOT / "results" / "DTI_Project_Final_Report_NEXUS_ATMS.md"

TEAM_MEMBERS = [
    ("Khushi Jangra", "S24CSEU0666"),
    ("Jaismeen Kaur", "S24CSEU0682"),
]
MENTOR_NAME = "Sushmita Das"


def add_h1(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)


def add_h2(doc: Document, title: str) -> None:
    doc.add_heading(title, level=2)


def add_h3(doc: Document, title: str) -> None:
    doc.add_heading(title, level=3)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build_doc() -> Document:
    doc = Document()

    doc.add_heading("NEXUS-ATMS: AI-Driven Urban Congestion Management and Adaptive Traffic Signal Control", 0)
    add_para(doc, "A PROJECT REPORT")
    add_para(doc, "Submitted by: Team 54")
    for name, enroll in TEAM_MEMBERS:
        add_para(doc, f"{name} ({enroll})")
    add_para(doc, f"Mentor: {MENTOR_NAME} Mam")
    add_para(doc, "School of Computer Science Engineering and Technology, Bennett University")
    add_para(doc, datetime.now().strftime("%B %Y"))

    add_h1(doc, "DECLARATION")
    add_para(
        doc,
        "We hereby declare that the work presented in this report, titled 'NEXUS-ATMS: AI-Driven Urban Congestion Management and Adaptive Traffic Signal Control', is an authentic record of our own work carried out at Bennett University during the academic session 2025-2026. The work presented has not been submitted elsewhere for the award of any degree.",
    )
    add_para(doc, "Signatures:")
    for name, enroll in TEAM_MEMBERS:
        add_para(doc, f"- {name} ({enroll})")

    add_h1(doc, "ACKNOWLEDGEMENT")
    add_para(
        doc,
        f"We express our sincere gratitude to our mentor, {MENTOR_NAME} Mam, and the faculty members of the School of Computer Science Engineering and Technology, Bennett University, for their guidance, feedback, and continuous support throughout this project. We also thank our peers for their constructive suggestions during design reviews and testing phases.",
    )

    add_h1(doc, "LIST OF TABLES")
    add_para(doc, "Table 1: Quantified Project Goals and Objectives")
    add_para(doc, "Table 2: Project Setup Decisions")
    add_para(doc, "Table 3: Stakeholders and Roles")
    add_para(doc, "Table 4: Project Resources")
    add_para(doc, "Table 5: Project Assumptions")
    add_para(doc, "Table 6: Communication Plan")
    add_para(doc, "Table 7: Deliverables")

    add_h1(doc, "LIST OF FIGURES")
    add_para(doc, "Figure 1: NEXUS-ATMS Layered Architecture")
    add_para(doc, "Figure 2: Use Case View")
    add_para(doc, "Figure 3: Class and Component View")
    add_para(doc, "Figure 4: Activity and Sequence Flow")
    add_para(doc, "Figure 5: Data Architecture")
    add_para(doc, "Figure 6: Dashboard User Interface")

    add_h1(doc, "LIST OF ABBREVIATIONS")
    add_para(doc, "ATMS - Adaptive Traffic Management System")
    add_para(doc, "RL - Reinforcement Learning")
    add_para(doc, "D3QN - Double Dueling Deep Q-Network")
    add_para(doc, "PPO - Proximal Policy Optimization")
    add_para(doc, "SUMO - Simulation of Urban MObility")
    add_para(doc, "TraCI - Traffic Control Interface")
    add_para(doc, "LSTM - Long Short-Term Memory")
    add_para(doc, "API - Application Programming Interface")

    add_h1(doc, "ABSTRACT")
    add_para(
        doc,
        "Urban congestion causes increased travel time, fuel wastage, emissions, and emergency response delays. This project presents NEXUS-ATMS, an AI-driven urban traffic management system that integrates reinforcement learning, prediction modules, anomaly detection, and a real-time dashboard. The system supports SUMO-based simulation and standalone operation. The reinforcement learning stack includes DQN/PPO agents and a custom D3QN pipeline with crash-safe checkpointing and exact resume capability. Supporting modules include LSTM-based traffic forecasting, anomaly detection, emergency corridor handling, and explainability tools for decision transparency. "
        "In evaluation, the RL agent substantially reduced average waiting time and queue length compared with fixed-timing control. On the primary benchmark, waiting time reduced by about 98.34% and queue length by about 90.66%, demonstrating strong congestion mitigation in the tested configuration. The platform is modular, supports report generation and demo packaging, and provides practical pathways for further scaling to multi-intersection control, robustness testing, and deployment-oriented validation.",
    )

    add_h1(doc, "1. INTRODUCTION")
    add_para(
        doc,
        "Intelligent transportation systems are shifting from static timing plans to adaptive control because traffic demand is non-stationary and context-sensitive. NEXUS-ATMS is designed as a practical AI platform that senses, predicts, and optimizes signal behavior in near real time while maintaining observability through dashboards and reports.",
    )

    add_h2(doc, "1.1 Problem Statement")
    add_para(
        doc,
        "Conventional fixed-time traffic signals cannot adapt to dynamic traffic fluctuations, resulting in unnecessary waiting, long queues, and inefficient corridor movement. The problem is to design an adaptive control framework that reduces congestion metrics while preserving operational reliability and reproducibility.",
    )

    add_h1(doc, "2. Background Research")
    add_para(
        doc,
        "Recent intelligent traffic control literature shows RL methods outperform fixed-time baselines in simulated environments, especially under variable demand. Value-based methods (DQN variants) offer direct action-value optimization, while policy-gradient methods (PPO) offer stable policy updates. Hybrid system design that combines predictive analytics with RL policy control is increasingly favored for operational robustness.",
    )

    add_h2(doc, "2.1 Proposed System")
    add_para(
        doc,
        "NEXUS-ATMS uses a layered architecture: ingestion (SUMO/vision/sensors), AI engine (RL + LSTM + anomaly), specialty modules (emergency, safety, carbon), backend (FastAPI + WebSocket), and presentation (authority dashboard + analytics). The core control loop continuously observes traffic state, selects signal actions, and updates policy performance through reward-driven feedback.",
    )

    add_h2(doc, "2.2 Goals and Objectives")
    add_para(doc, "Table 1: Quantified Project Goals and Objectives")
    add_para(doc, "1) Reduce average waiting time by more than 50% relative to fixed timing.")
    add_para(doc, "2) Reduce average queue length by more than 40% relative to fixed timing.")
    add_para(doc, "3) Provide crash-safe model checkpointing and resume support.")
    add_para(doc, "4) Provide explainable and visual analytics for stakeholders.")
    add_para(doc, "5) Deliver a reproducible project package for evaluation and demo submission.")

    add_h1(doc, "3. Project Planning")
    add_h2(doc, "3.1 Project Lifecycle")
    add_para(
        doc,
        "The project followed an iterative lifecycle: requirement capture, architecture design, environment setup, baseline development, RL integration, evaluation, stabilization, and demo/report packaging. Each phase ended with validation checkpoints and artifact generation.",
    )

    add_h2(doc, "3.2 Project Setup")
    add_para(doc, "Table 2: Project Setup Decisions")
    add_para(doc, "- Language/Runtime: Python 3.13")
    add_para(doc, "- Simulation: SUMO + TraCI")
    add_para(doc, "- ML stack: PyTorch, Stable-Baselines3")
    add_para(doc, "- API/Dashboard: FastAPI + HTML/JS frontend")
    add_para(doc, "- Logging/Artifacts: JSON, model checkpoints, evaluation charts")

    add_h2(doc, "3.3 Stakeholders")
    add_para(doc, "Table 3: Stakeholders and Roles")
    add_para(doc, "- Mentor: Technical guidance and review")
    add_para(doc, "- Team Members: Development, testing, documentation, demo")
    add_para(doc, "- Faculty/Evaluators: Academic assessment")
    add_para(doc, "- End Users (conceptual): Traffic authority and citizens")

    add_h2(doc, "3.4 Project Resources")
    add_para(doc, "Table 4: Project Resources")
    add_para(doc, "- Hardware: Ryzen 5 5600H, RTX 2050 (4 GB), 16 GB RAM")
    add_para(doc, "- Software: SUMO, PyTorch, SB3, FastAPI")
    add_para(doc, "- Data: Simulated traffic flows and generated scenarios")

    add_h2(doc, "3.5 Assumptions")
    add_para(doc, "Table 5: Project Assumptions")
    add_para(doc, "- Simulation behavior is representative for comparative evaluation.")
    add_para(doc, "- Benchmark metrics are measured consistently across baseline and RL runs.")
    add_para(doc, "- GPU availability improves speed but is not mandatory for correctness.")

    add_h1(doc, "4. Project Tracking")
    add_h2(doc, "4.1 Tracking")
    add_para(
        doc,
        "Project tracking used version-controlled source artifacts, structured logs, checkpointed model folders, and generated benchmark summaries. Training and evaluation outputs were retained under logs/, models/, and results/ for traceability.",
    )

    add_h2(doc, "4.2 Communication Plan")
    add_para(doc, "Table 6: Communication Plan")
    add_para(doc, "- Daily standup style progress sync")
    add_para(doc, "- Weekly mentor review and corrective actions")
    add_para(doc, "- Milestone reports for architecture, training, and evaluation")

    add_h2(doc, "4.3 Deliverables")
    add_para(doc, "Table 7: Deliverables")
    add_para(doc, "- Adaptive traffic control codebase")
    add_para(doc, "- Trained model artifacts and checkpoints")
    add_para(doc, "- Evaluation results and charts")
    add_para(doc, "- Dashboard and demo package")
    add_para(doc, "- Final report and references")

    add_h1(doc, "5. SYSTEM ANALYSIS AND DESIGN")
    add_h2(doc, "5.1 Overall Description")
    add_para(
        doc,
        "NEXUS-ATMS is a modular AI traffic platform with independent but interoperable components. It supports sensing, policy inference, prediction, anomaly monitoring, and operator-facing visualization. The architecture is intentionally layered to simplify maintenance and expansion.",
    )

    add_h2(doc, "5.2 Users and Roles")
    add_para(doc, "- Traffic Authority: Monitor KPIs, apply overrides, evaluate incidents.")
    add_para(doc, "- Analyst/Researcher: Inspect training curves and benchmark outcomes.")
    add_para(doc, "- System Admin: Manage deployment, logs, and model updates.")

    add_h2(doc, "5.3 Design diagrams/Architecture/ UML diagrams/ Flow Charts/ E-R diagrams")
    add_h3(doc, "5.3.1 Product Backlog Items")
    add_para(doc, "Adaptive control loop, checkpoint resume, evaluation pipeline, plotting, dashboard integration, and demo bundle generation were prioritized backlog items.")

    add_h3(doc, "5.3.2 Architecture Diagram")
    add_para(doc, "Layered architecture: ingestion -> AI engine -> specialty modules -> backend APIs -> dashboard.")

    add_h3(doc, "5.3.3 Use Case Diagram")
    add_para(doc, "Primary use cases include run simulation, train agent, evaluate policy, monitor dashboard, and generate reports.")

    add_h3(doc, "5.3.4 Class Diagram")
    add_para(doc, "Core classes include environment abstractions, RL agent classes, replay/checkpoint utilities, and reporting modules.")

    add_h3(doc, "5.3.5 Activity Diagrams")
    add_para(doc, "Main activity flow: initialize environment -> observe state -> select action -> step environment -> compute reward -> update model -> log and checkpoint.")

    add_h3(doc, "5.3.6 Sequence Diagram")
    add_para(doc, "Client triggers train/evaluate -> backend initializes components -> agent/environment interaction loop -> artifacts saved -> dashboard renders status.")

    add_h3(doc, "5.3.7 Data Architecture")
    add_para(doc, "Operational data stored as JSON artifacts, model binaries, and log directories, enabling reproducible experiments and post-run analysis.")

    add_h1(doc, "6. User Interface")
    add_h2(doc, "6.1 UI Description")
    add_para(doc, "The dashboard provides a control-centric authority view, analytics panels, and system health indicators through API and WebSocket updates.")

    add_h2(doc, "6.2 UI Mockup")
    add_para(doc, "The UI includes traffic status, control cards, KPI tiles, model insights, and chart panels suitable for live demonstration and decision support.")

    add_h1(doc, "7. Algorithms/Pseudo Code OF CORE FUNCTIONALITY")
    add_para(doc, "Pseudo Code: Adaptive RL Control Loop")
    add_para(doc, "1) Initialize environment E and policy network Q")
    add_para(doc, "2) For each timestep t: observe state s_t")
    add_para(doc, "3) Select action a_t using epsilon-greedy/Q policy")
    add_para(doc, "4) Execute action in E, receive r_t and s_(t+1)")
    add_para(doc, "5) Store transition in replay buffer")
    add_para(doc, "6) Sample mini-batch, compute TD target, update Q")
    add_para(doc, "7) Periodically sync target network, evaluate, and checkpoint")
    add_para(doc, "8) Continue until convergence or time budget")

    add_h1(doc, "8. Project Closure")
    add_h2(doc, "8.1 Goals / Vision")
    add_para(doc, "The project vision is an adaptive, explainable, and reproducible traffic management stack that reduces congestion while remaining practical for demonstration and future deployment studies.")

    add_h2(doc, "8.2 Delivered Solution")
    add_para(doc, "Delivered components include RL training/evaluation pipelines, D3QN resume-safe checkpointing, analytics plotting, benchmark reports, and a demo submission bundle generator.")
    add_para(doc, "Observed benchmark results: waiting time reduction ~98.34%, queue reduction ~90.66% versus fixed timing in the evaluated setup.")

    add_h2(doc, "8.3 Remaining Work")
    add_para(doc, "Remaining work includes larger-scale multi-intersection validation, stronger robustness tests, expanded real-world sensor integration, and production-grade fault-tolerance hardening for long-duration operations.")

    add_h1(doc, "REFERENCES")
    add_para(doc, "[1] SUMO - Simulation of Urban MObility, Eclipse Foundation.")
    add_para(doc, "[2] Mnih et al., Human-level control through deep reinforcement learning, Nature, 2015.")
    add_para(doc, "[3] Schulman et al., Proximal Policy Optimization Algorithms, arXiv, 2017.")
    add_para(doc, "[4] PyTorch Documentation, https://pytorch.org")
    add_para(doc, "[5] Stable-Baselines3 Documentation, https://stable-baselines3.readthedocs.io")
    add_para(doc, "[6] FastAPI Documentation, https://fastapi.tiangolo.com")

    return doc


def build_markdown() -> str:
    return """# NEXUS-ATMS: AI-Driven Urban Congestion Management and Adaptive Traffic Signal Control\n\nA full report aligned with the DTI final report template has been generated as a Word document:\n\n- `results/DTI_Project_Final_Report_NEXUS_ATMS.docx`\n\nThis markdown companion is included for versioning convenience.\n"""


def main() -> int:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

    doc = build_doc()
    doc.save(OUT_DOCX)

    OUT_MD.write_text(build_markdown(), encoding="utf-8")

    print(f"[OK] Report generated: {OUT_DOCX}")
    print(f"[OK] Companion note: {OUT_MD}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
